# -*- coding: utf-8 -*-
"""
视频流分析服务（整合版）
- MQ 主流程：analyze_from_stream，消费 RabbitMQ 任务，YOLO + 视觉模型 + RAG 规章制度查询
- HTTP 接口：analyze_url，供 /upload-video 直接调用
原 video_analyzer.py 已废弃，本文件为唯一分析入口。
"""
import os
import json
import logging
import cv2
from typing import Dict, Any, List, Optional
from concurrent.futures import ThreadPoolExecutor

from openai import OpenAI
from docx import Document as DocxDocument
from langchain.chat_models import init_chat_model
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings
from app.models.schemas import VideoTaskMessage, VideoTaskResult, EventInfo

logger = logging.getLogger(__name__)

# ─────────────────────────── Prompts ───────────────────────────

# 整合自 video_analyzer.py：含 JSON 示例、start_second/end_second、明确禁止 markdown
ANALYSIS_PROMPT = '''
你是一位高度专业的铜矿生产视频行为分析专家，同时具备严谨的数据处理能力。
    你的任务是检查视频中是否有违规行为，同时读取叠加在帧图像上的元数据（日期、时间、用户编号、单位编号、序列号），并在违规行为开始到结束期间记录这些元数据。
    任务目标：
    1.视频帧与元数据解析：
        视频提取：处理每一帧视频图像，并同时读取叠加在图像上的文本信息。
        元数据提取：精准提取以下五项元数据。
            日期：视频拍摄的日期。
            时间：视频拍摄的具体时间。
            用户编号：视频中的用户编号。
            单位编号：视频中的单位编号。
            序列号：视频中的序列号。
    2.事件记录与关联元数据：
        对每个违规行为，记录事件的「开始时间」和「结束时间」,同时关联该行为中涉及的用户编号、单位编号、序列号以及日期。
    3.输出要求：
        将违规行为记录为可被机器直接解析的 JSON 数组，无需任何 Markdown 代码块标记。
        若视频中未识别到任何事件，则每个字段都返回无。
        每个 JSON 对象代表一个事件，必须包含以下字段：
            event_description： 对违规行为的详细说明。
            date： 事件发生日期（YYYY-MM-DD）。
            start_time： 事件发生时间（HH:MM:SS）。
            end_time： 事件结束时间（HH:MM:SS）。
            user_number： 涉事用户的编号。
            unit_number： 涉事单位的编号。
            serial_number： 相关的设备或视频序列号。
            start_second: 事件在视频第几秒开始。
            end_second: 事件在视频第几秒结束。
        JSON 对象结构示例如下：
            [
                {
                "event_description": "",
                "date": "",
                "start_time": "",
                "end_time": "",
                "user_number": "",
                "unit_number": "",
                "serial_number": ""，
                "start_second": "",
                "end_second": ""
                }
            ]
'''

RAG_PROMPT = ChatPromptTemplate.from_template(
    """
    请根据下面提供的上下文信息来回答问题。
    请确保你的回答完全基于这些上下文。
    如果上下文中没有足够的信息来回答问题，请直接告知："抱歉，我无法根据提供的上下文找到相关信息来回答此问题。"

    上下文:
    {context}

    问题: {question}

    回答:
    """
)

# 线程池：用于执行阻塞的CPU/GPU密集型操作
_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix="video_analyzer")


class StreamVideoAnalyzer:
    """
    流式视频分析器（整合版，唯一分析入口）
    - MQ 流程：analyze_from_stream
    - HTTP 流程：analyze_url
    """

    def __init__(self):
        self.settings = get_settings()
        self._vision_client: Optional[OpenAI] = None   # 原生 OpenAI 客户端，支持 temperature 等参数
        self._thinking_model = None                     # LangChain 模型，用于 RAG 规章制度查询
        self._vector_store: Optional[InMemoryVectorStore] = None
        self._vector_store_failed = False

    @property
    def vision_client(self) -> OpenAI:
        """懒加载视觉模型客户端（原生 openai，可传 temperature/top_p/presence_penalty 等）"""
        if self._vision_client is None:
            self._vision_client = OpenAI(
                base_url=self.settings.vision_model_url,
                api_key=self.settings.model_api_key,
                timeout=self.settings.model_timeout,
            )
        return self._vision_client

    @property
    def thinking_model(self):
        """懒加载 RAG 回答模型（LangChain）"""
        if self._thinking_model is None:
            self._thinking_model = init_chat_model(
                model=self.settings.thinking_model_name,
                model_provider="openai",
                base_url=self.settings.thinking_model_url,
                api_key=self.settings.model_api_key,
                timeout=self.settings.model_timeout,
            )
        return self._thinking_model

    def _strip_thinking_content(self, content: str) -> str:
        """去掉模型思考内容，只保留思维链标记之后的输出结果"""
        for end_tag in ("</think>", "<|im_end|>"):
            if end_tag in content:
                content = content.split(end_tag, 1)[1]
                break
        return content.strip()

    def _resolve_regulations_docx_path(self) -> str:
        """解析规章制度 docx 路径（支持相对项目根目录）"""
        docx_path = self.settings.regulations_docx_path
        if os.path.isabs(docx_path):
            return docx_path
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        return os.path.join(project_root, docx_path)

    def _load_regulation_documents(self) -> List[Document]:
        """从 docx 文件加载规章制度文本"""
        docx_path = self._resolve_regulations_docx_path()
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"规章制度文件不存在: {docx_path}")

        docx = DocxDocument(docx_path)
        paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return [Document(page_content="\n".join(paragraphs))]

    def _get_vector_store(self) -> Optional[InMemoryVectorStore]:
        """懒加载规章制度向量库"""
        if self._vector_store_failed:
            return None
        if self._vector_store is not None:
            return self._vector_store

        try:
            docs = self._load_regulation_documents()
            text_splitter = RecursiveCharacterTextSplitter()
            texts = text_splitter.split_documents(docs)
            embeddings = OpenAIEmbeddings(
                model=self.settings.embedding_model_name,
                base_url=self.settings.embedding_base_url,
                api_key=self.settings.embedding_api_key or self.settings.model_api_key,
            )
            vector_store = InMemoryVectorStore(embeddings)
            vector_store.add_documents(texts)
            self._vector_store = vector_store
            logger.info(f"[RAG] 向量库初始化完成，文档块数={len(texts)}")
            return self._vector_store
        except Exception as e:
            self._vector_store_failed = True
            logger.error(f"[RAG] 向量库初始化失败: {e}")
            return None

    def _call_vision_model(self, video_url: str, extra_text: str = "") -> str:
        """
        调用视觉模型，返回原始文本。
        使用原生 OpenAI 客户端，传入与 video_analyzer.py 一致的推理参数。
        """
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "video_url", "video_url": {"url": video_url}},
                    {"type": "text", "text": ANALYSIS_PROMPT + extra_text},
                ],
            }
        ]
        response = self.vision_client.chat.completions.create(
            model=self.settings.vision_model_name,
            messages=messages,
            max_tokens=4096,
            temperature=1.0,
            top_p=0.95,
            presence_penalty=1.5,
            extra_body={"top_k": 20},
        )
        return response.choices[0].message.content or ""
    
    def _run_yolo_detection(self, video_url: str) -> List[str]:
        """
        本地直接运行 YOLO 流式检测，不依赖 LLM tool calling。
        返回逐秒检测结果列表，供后续拼入视觉模型 prompt。
        """
        settings = self.settings
        logger.info(f"[YOLO] 开始流式读取视频: {video_url[:100]}...")

        if not os.path.exists(settings.yolo_model_path):
            logger.warning(f"YOLO模型不存在: {settings.yolo_model_path}，跳过检测")
            return []

        try:
            from ultralytics import YOLO
            model = YOLO(settings.yolo_model_path)

            cap = cv2.VideoCapture(video_url)
            if not cap.isOpened():
                logger.error(f"无法打开视频流: {video_url[:100]}...")
                return []

            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_interval = max(int(fps), 1) if fps > 0 else 30

            detection_results: List[str] = []
            frame_count = 0
            second_count = 0
            max_seconds = 60

            while cap.isOpened() and second_count < max_seconds:
                success, frame = cap.read()
                if not success:
                    break

                if frame_count % frame_interval == 0:
                    second_count += 1
                    results = model(frame, verbose=False)
                    objects = [
                        results[0].names[int(box.cls[0])]
                        for box in results[0].boxes
                    ]
                    if objects:
                        detection_results.append(
                            f"第{second_count}秒检测到: {', '.join(set(objects))}"
                        )

                frame_count += 1

            cap.release()
            logger.info(f"[YOLO] 流式检测完成，处理了 {second_count} 秒视频，共 {len(detection_results)} 条结果")
            return detection_results

        except Exception as e:
            logger.error(f"YOLO检测失败: {e}")
            return []
    
    def analyze_from_stream(self, task: VideoTaskMessage) -> VideoTaskResult:
        """
        从预签名URL流式分析视频
        
        Args:
            task: 视频任务消息
            
        Returns:
            VideoTaskResult: 处理结果
        """
        import time
        start_time = time.time()
        
        logger.info(f"[分析] 开始处理任务: taskId={task.task_id}, videoId={task.video_id}")
        logger.info(f"[分析] 预签名URL: {task.presigned_url[:80]}...")
        
        try:
            # 第一步：本地 YOLO 检测（不依赖 LLM tool calling）
            yolo_results = self._run_yolo_detection(task.presigned_url)
            yolo_summary = (
                "\n\n【YOLO辅助检测结果】\n" + "\n".join(yolo_results)
                if yolo_results else ""
            )

            # 第二步：直接调用视觉模型（不使用 agent，避免 tool_choice="auto" 报错）
            vision_messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "video_url", "video_url": {"url": task.presigned_url}},
                        {"type": "text", "text": ANALYSIS_PROMPT + yolo_summary}
                    ]
                }
            ]

            content = self._call_vision_model(task.presigned_url, yolo_summary)

            # 解析事件
            events_data = self._parse_events(content)

            # RAG 规章制度查询（thinking 模型不可用时降级：直接把 vision 模型识别的事件当 unsafe_events）
            unsafe_events = self._analyze_safety_sync(events_data, fallback_events=events_data)
            
            process_time = time.time() - start_time
            logger.info(f"[分析] 任务完成: taskId={task.task_id}, 耗时={process_time:.2f}s, 事件数={len(events_data)}")
            
            # 如果有违规事件，尝试捕获违规帧，并根据视频时长校验违规时间
            violation_frame = None
            violation_timestamp = None
            if unsafe_events:
                try:
                    cap = cv2.VideoCapture(task.presigned_url)
                    if cap.isOpened():
                        # 获取视频时长，用于校验违规时间是否合理（防止模型误用画面上的时钟时间）
                        frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
                        fps = cap.get(cv2.CAP_PROP_FPS)
                        video_duration = (frame_count / fps) if fps and fps > 0 and frame_count > 0 else 0.0
                        if video_duration > 0:
                            self._sanitize_violation_times(unsafe_events, video_duration)
                        
                        # 尝试从第一个违规事件中提取时间戳并捕获违规帧
                        first_unsafe = unsafe_events[0]
                        violation_timestamp = first_unsafe.get('start_second')
                        if violation_timestamp is not None and fps and fps > 0:
                            frame_number = int(violation_timestamp * fps)
                            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
                            success_read, violation_frame = cap.read()
                            if not success_read:
                                violation_frame = None
                    cap.release()
                except Exception as e:
                    logger.warning(f"捕获违规帧失败: {e}")
            
            return VideoTaskResult(
                task_id=task.task_id,
                video_id=task.video_id,
                success=True,
                events=[EventInfo(**e) for e in events_data] if events_data else None,
                unsafe_events=[EventInfo(**e) for e in unsafe_events] if unsafe_events else None,
                process_time=process_time,
                violation_frame=violation_frame,
                violation_timestamp=violation_timestamp,
                raw_analysis=content  # 保留原始文本，供 ai_description 兜底
            )

        except Exception as e:
            process_time = time.time() - start_time
            logger.error(f"[分析] 任务失败: taskId={task.task_id}, error={e}")

            return VideoTaskResult(
                task_id=task.task_id,
                video_id=task.video_id,
                success=False,
                error_message=str(e),
                process_time=process_time
            )

    def analyze_url(self, video_url: str) -> Dict[str, Any]:
        """
        HTTP 接口入口（替代已废弃的 VideoAnalyzer.analyze）。
        直接调用视觉模型分析指定 URL 的视频，返回事件列表字典。
        不包含 YOLO 检测和安全合规二次审查（HTTP 上传场景通常不需要）。
        """
        import time as _time
        logger.info(f"[analyze_url] 开始分析: {video_url[:80]}...")
        t0 = _time.time()
        content = self._call_vision_model(video_url)
        events_data = self._parse_events(content)
        elapsed = _time.time() - t0
        logger.info(f"[analyze_url] 分析完成，耗时={elapsed:.2f}s，事件数={len(events_data)}")
        return {
            "success": True,
            "events": events_data,
            "total_events": len(events_data),
        }

    def _parse_events(self, content: str) -> List[Dict]:
        """
        从模型回复中提取事件 JSON 列表。
        依次尝试：
          1. 去掉 markdown 代码块后直接 json.loads
          2. 从文本中定位第一个 '[' 到最后一个 ']' 的片段
          3. 从文本中定位第一个 '{' 到最后一个 '}' 的片段（单对象包成列表）
        全部失败时记录原始内容供排查，返回空列表。
        """
        # 记录原始响应（前 500 字符）便于排查
        preview = content[:500].replace('\n', ' ') if content else ''
        logger.info(f"[parse_events] 模型原始回复(前500字): {preview}")

        if not content or not content.strip():
            logger.warning("[parse_events] 模型返回空内容")
            return []

        cleaned = content.strip()

        # 0. 剥离思维链：hrylora 等推理模型会在 </think> 前输出推理过程，真正结果在其后
        for end_tag in ("</think>", "<|im_end|>"):
            if end_tag in cleaned:
                cleaned = cleaned.split(end_tag, 1)[1].strip()
                logger.info(f"[parse_events] 检测到思维链标记 {end_tag!r}，已剥离，剩余内容: {cleaned[:200]}")
                break

        # 1. 去掉 markdown 代码块标记
        if cleaned.startswith("```"):
            lines = cleaned.split('\n')
            # 去掉首行（```json 或 ```）和末尾的 ```
            inner_lines = lines[1:]
            if inner_lines and inner_lines[-1].strip() == "```":
                inner_lines = inner_lines[:-1]
            cleaned = '\n'.join(inner_lines).strip()

        # 2. 直接尝试解析
        try:
            data = json.loads(cleaned)
            return self._normalize_events(data)
        except json.JSONDecodeError:
            pass

        # 3. 提取第一个 '[' … 最后一个 ']' 片段
        start = cleaned.find('[')
        end = cleaned.rfind(']')
        if start != -1 and end != -1 and end > start:
            try:
                data = json.loads(cleaned[start:end + 1])
                return self._normalize_events(data)
            except json.JSONDecodeError:
                pass

        # 4. 提取第一个 '{' … 最后一个 '}' 片段（单对象）
        start = cleaned.find('{')
        end = cleaned.rfind('}')
        if start != -1 and end != -1 and end > start:
            try:
                data = json.loads(cleaned[start:end + 1])
                return self._normalize_events(data)
            except json.JSONDecodeError:
                pass

        # 全部失败：打 WARNING 并记录原始内容
        logger.warning(f"[parse_events] 无法解析JSON，原始回复: {content[:300]}")
        return []

    def _normalize_events(self, data) -> List[Dict]:
        """将解析结果统一为列表格式"""
        if isinstance(data, list):
            return [item for item in data if isinstance(item, dict)]
        if isinstance(data, dict):
            # 模型有时返回 {"events": [...]} 包装
            for key in ("events", "结果", "事件列表", "event_list"):
                if key in data and isinstance(data[key], list):
                    return [item for item in data[key] if isinstance(item, dict)]
            return [data]
        return []
    
    def _analyze_safety_sync(self, events: List[Dict],
                             fallback_events: Optional[List[Dict]] = None) -> List[Dict]:
        """
        同步 RAG 规章制度查询。
        对每个违规事件检索相关制度并补充 regulations 字段；
        若 RAG 不可用则降级为直接使用 vision 模型的事件描述。
        """
        vector_store = self._get_vector_store()
        if vector_store is None:
            return self._fallback_unsafe_events(fallback_events or events)

        unsafe_events = []
        for event in events:
            description = event.get('event_description', event.get('事件描述', ''))
            if not description or description == "无":
                continue

            try:
                question = (
                    f"针对以下违规行为，请给出相关的安全事故隐患排查治理规章制度：{description}"
                )
                retrieved_docs = vector_store.similarity_search(
                    question, k=self.settings.rag_top_k
                )
                docs_content = "\n\n".join(doc.page_content for doc in retrieved_docs)
                answer = self.thinking_model.invoke(
                    RAG_PROMPT.format(question=question, context=docs_content)
                )
                regulations_text = self._strip_thinking_content(answer.content)

                start_time_str = event.get('start_time', event.get('开始时间', ''))
                end_time_str = event.get('end_time', event.get('结束时间', ''))
                unsafe_events.append({
                    "event_description": description,
                    "date": event.get('date', event.get('日期', '')),
                    "start_time": start_time_str,
                    "end_time": end_time_str,
                    "start_second": event.get('start_second') or self._parse_time_to_seconds(start_time_str),
                    "end_second": event.get('end_second') or self._parse_time_to_seconds(end_time_str),
                    "user_number": event.get('user_number', event.get('用户编号', '')),
                    "unit_number": event.get('unit_number', event.get('单位编号', '')),
                    "serial_number": event.get('serial_number', event.get('序列号', '')),
                    "regulations": regulations_text,
                })
                logger.info(f"[RAG] 违规事件: {description[:80]}...")
            except Exception as e:
                logger.warning(f"[RAG] 规章制度查询失败: {e}")

        if not unsafe_events and fallback_events:
            logger.warning("[RAG] 查询无结果或全部失败，降级为 vision 事件描述作为 unsafe_events")
            return self._fallback_unsafe_events(fallback_events)

        return unsafe_events

    def _fallback_unsafe_events(self, events: List[Dict]) -> List[Dict]:
        """RAG 不可用时，使用 vision 模型原始事件作为 unsafe_events"""
        unsafe_events = []
        for ev in events:
            desc = ev.get('event_description', ev.get('事件描述', ''))
            if not desc or desc == "无":
                continue
            unsafe_events.append({
                "event_description": desc,
                "date": ev.get('date', ev.get('日期', '')),
                "start_time": ev.get('start_time', ev.get('开始时间', '')),
                "end_time": ev.get('end_time', ev.get('结束时间', '')),
                "start_second": ev.get('start_second'),
                "end_second": ev.get('end_second'),
                "user_number": ev.get('user_number', ev.get('用户编号', '')),
                "unit_number": ev.get('unit_number', ev.get('单位编号', '')),
                "serial_number": ev.get('serial_number', ev.get('序列号', '')),
                "regulations": "",
            })
        return unsafe_events
    
    def _sanitize_violation_times(self, unsafe_events: List[Dict], video_duration: float) -> None:
        """
        根据视频时长校验违规时间，超出范围则置为 None。
        防止模型误用画面上的录制时钟时间（如 18:18:42）导致错误。
        """
        for ev in unsafe_events:
            start_sec = ev.get('start_second')
            end_sec = ev.get('end_second')
            if start_sec is not None and start_sec > video_duration:
                logger.warning(f"违规时间 start_second={start_sec}s 超出视频时长 {video_duration:.1f}s，已置空")
                ev['start_second'] = None
            if end_sec is not None and end_sec > video_duration:
                logger.warning(f"违规时间 end_second={end_sec}s 超出视频时长 {video_duration:.1f}s，已置空")
                ev['end_second'] = None
    
    def _parse_time_to_seconds(self, time_str: str) -> Optional[float]:
        """
        将时间字符串解析为秒数
        支持格式: HH:MM:SS, MM:SS, SS
        """
        if not time_str:
            return None
        
        try:
            parts = time_str.strip().split(':')
            if len(parts) == 3:  # HH:MM:SS
                return int(parts[0]) * 3600 + int(parts[1]) * 60 + float(parts[2])
            elif len(parts) == 2:  # MM:SS
                return int(parts[0]) * 60 + float(parts[1])
            elif len(parts) == 1:  # SS
                return float(parts[0])
            return None
        except (ValueError, IndexError):
            return None


# 全局分析器实例
_analyzer: Optional[StreamVideoAnalyzer] = None


def get_stream_analyzer() -> StreamVideoAnalyzer:
    """获取流式分析器单例"""
    global _analyzer
    if _analyzer is None:
        _analyzer = StreamVideoAnalyzer()
    return _analyzer


def get_executor() -> ThreadPoolExecutor:
    """获取线程池"""
    return _executor
